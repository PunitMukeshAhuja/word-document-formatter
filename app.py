
import io
import re
import hashlib
import html
import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

st.set_page_config(
    page_title="Word Document Formatter",
    page_icon="📝",
    layout="wide"
)

st.title("📝 Word Document Formatter")
st.caption(
    "Upload a Word document, review its detected structure and proposed formatting, "
    "preview the document visually, then download the formatted Word file."
)

CLASS_OPTIONS = ["Title", "Heading 1", "Heading 2", "Heading 3", "Body"]

# -------------------------------------------------------
# Helpers
# -------------------------------------------------------

def file_signature(file_bytes):
    return hashlib.md5(file_bytes).hexdigest()

def paragraph_text(paragraph):
    return paragraph.text.strip()

def average_font_size(paragraph):
    sizes = [
        run.font.size.pt
        for run in paragraph.runs
        if run.font.size
    ]
    return sum(sizes) / len(sizes) if sizes else None

def mostly_bold(paragraph):
    runs = [r for r in paragraph.runs if r.text.strip()]
    if not runs:
        return False
    bold_runs = sum(1 for r in runs if r.bold is True)
    return bold_runs / len(runs) >= 0.6

def is_all_caps(text):
    letters = [c for c in text if c.isalpha()]
    return bool(letters) and all(c.isupper() for c in letters)

def detect_heading_level(paragraph, auto_detect, body_size):
    text = paragraph_text(paragraph)

    if not text:
        return None

    style_name = paragraph.style.name if paragraph.style else ""

    # Existing Word styles always take priority.
    if style_name == "Title":
        return 0
    if style_name == "Heading 1":
        return 1
    if style_name == "Heading 2":
        return 2
    if style_name == "Heading 3":
        return 3

    # With auto-detection OFF, do not infer structure.
    if not auto_detect:
        return None

    if len(text) > 120:
        return None

    # 1. Introduction / 1.1 Scope / 1.1.1 Detail
    numbered = re.match(r"^\s*(\d+(?:\.\d+){0,2})[\.\)]?\s+\S+", text)
    if numbered:
        depth = numbered.group(1).count(".") + 1
        return min(depth, 3)

    avg_size = average_font_size(paragraph)
    bold = mostly_bold(paragraph)
    caps = is_all_caps(text)

    if (
        len(text) <= 70
        and caps
        and bold
        and avg_size is not None
        and avg_size >= body_size + 4
    ):
        return 0

    if (
        len(text) <= 60
        and caps
        and (bold or (avg_size is not None and avg_size >= body_size + 2))
    ):
        return 1

    if len(text) <= 70 and bold:
        if avg_size is not None and avg_size >= body_size + 3:
            return 1
        if avg_size is not None and avg_size >= body_size + 1:
            return 2
        return 2

    if len(text) <= 80 and avg_size is not None:
        if avg_size >= body_size + 4:
            return 1
        if avg_size >= body_size + 2:
            return 2
        if avg_size >= body_size + 1:
            return 3

    return None

def classification_label(level):
    return {
        0: "Title",
        1: "Heading 1",
        2: "Heading 2",
        3: "Heading 3",
        None: "Body"
    }[level]

def format_properties(
    final_type,
    body_size,
    title_size,
    heading1_size,
    heading2_size,
    heading3_size,
    body_alignment,
    line_spacing,
    space_after
):
    if final_type == "Title":
        return {
            "size": title_size,
            "bold": True,
            "alignment": "Center",
            "space_before": 0,
            "space_after": 12,
            "line_spacing": line_spacing
        }

    if final_type == "Heading 1":
        return {
            "size": heading1_size,
            "bold": True,
            "alignment": "Left",
            "space_before": 10,
            "space_after": 6,
            "line_spacing": line_spacing
        }

    if final_type == "Heading 2":
        return {
            "size": heading2_size,
            "bold": True,
            "alignment": "Left",
            "space_before": 8,
            "space_after": 4,
            "line_spacing": line_spacing
        }

    if final_type == "Heading 3":
        return {
            "size": heading3_size,
            "bold": True,
            "alignment": "Left",
            "space_before": 6,
            "space_after": 3,
            "line_spacing": line_spacing
        }

    return {
        "size": body_size,
        "bold": False,
        "alignment": body_alignment,
        "space_before": 0,
        "space_after": space_after,
        "line_spacing": line_spacing
    }

def build_base_preview(file_bytes, auto_detect, body_size):
    doc = Document(io.BytesIO(file_bytes))
    rows = []

    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph_text(paragraph)
        if not text:
            continue

        detected = classification_label(
            detect_heading_level(paragraph, auto_detect, body_size)
        )

        avg_size = average_font_size(paragraph)

        rows.append({
            "Paragraph index": idx,
            "Detected as": detected,
            "Final type": detected,
            "Text preview": text[:180] + ("..." if len(text) > 180 else ""),
            "Original style": paragraph.style.name if paragraph.style else "",
            "Original size": round(avg_size, 1) if avg_size else None,
            "Originally bold": "Yes" if mostly_bold(paragraph) else "No"
        })

    return pd.DataFrame(rows)

def apply_saved_overrides(df, overrides):
    if df.empty:
        return df

    result = df.copy()
    result["Final type"] = result.apply(
        lambda row: overrides.get(
            int(row["Paragraph index"]),
            row["Detected as"]
        ),
        axis=1
    )
    return result

def add_output_columns(
    df,
    body_size,
    title_size,
    heading1_size,
    heading2_size,
    heading3_size,
    body_alignment,
    line_spacing,
    space_after
):
    if df.empty:
        return df

    result = df.copy()

    props = result["Final type"].apply(
        lambda final_type: format_properties(
            final_type,
            body_size,
            title_size,
            heading1_size,
            heading2_size,
            heading3_size,
            body_alignment,
            line_spacing,
            space_after
        )
    )

    result["Output size"] = props.apply(lambda p: p["size"])
    result["Output bold"] = props.apply(lambda p: "Yes" if p["bold"] else "No")
    result["Output alignment"] = props.apply(lambda p: p["alignment"])
    result["Line spacing"] = props.apply(lambda p: p["line_spacing"])
    result["Space after"] = props.apply(lambda p: p["space_after"])

    return result

def set_run_font(run, name, size, bold=None):
    run.font.name = name
    run.font.size = Pt(size)

    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts

    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)

    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)

    if bold is not None:
        run.bold = bold

def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()

    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")

    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"

    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")

    run._r.append(fld_char1)
    run._r.append(instr_text)
    run._r.append(fld_char2)

def build_final_type_map(final_df):
    return {
        int(row["Paragraph index"]): row["Final type"]
        for _, row in final_df.iterrows()
    }

def render_print_preview(
    file_bytes,
    final_type_map,
    font_name,
    body_size,
    title_size,
    heading1_size,
    heading2_size,
    heading3_size,
    body_alignment,
    line_spacing,
    space_after,
    margin
):
    doc = Document(io.BytesIO(file_bytes))

    paragraphs_html = []

    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph_text(paragraph)
        if not text:
            continue

        final_type = final_type_map.get(idx, "Body")

        props = format_properties(
            final_type,
            body_size,
            title_size,
            heading1_size,
            heading2_size,
            heading3_size,
            body_alignment,
            line_spacing,
            space_after
        )

        align_css = {
            "Left": "left",
            "Center": "center",
            "Right": "right",
            "Justified": "justify"
        }[props["alignment"]]

        weight = "700" if props["bold"] else "400"

        escaped = html.escape(text)

        paragraphs_html.append(
            f"""
            <div style="
                font-family:{html.escape(font_name)};
                font-size:{props['size']}pt;
                font-weight:{weight};
                text-align:{align_css};
                line-height:{props['line_spacing']};
                margin-top:{props['space_before']}pt;
                margin-bottom:{props['space_after']}pt;
                overflow-wrap:anywhere;
            ">{escaped}</div>
            """
        )

    margin_px = int(float(margin) * 96)

    content = "\n".join(paragraphs_html)

    return f"""
    <div style="
        background:#e9edf2;
        padding:24px;
        border-radius:10px;
        overflow:auto;
        max-height:900px;
    ">
        <div style="
            width:min(794px, 100%);
            min-height:1123px;
            box-sizing:border-box;
            margin:0 auto 24px auto;
            background:white;
            padding:{margin_px}px;
            box-shadow:0 2px 14px rgba(0,0,0,0.18);
            color:#111;
        ">
            {content}
        </div>
    </div>
    """

def format_table(table, font_name, body_size):
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.line_spacing = 1.0
                paragraph.paragraph_format.space_after = Pt(0)

                for run in paragraph.runs:
                    set_run_font(run, font_name, body_size)

def format_document(
    file_bytes,
    final_type_map,
    font_name,
    body_size,
    line_spacing,
    alignment_enum,
    body_alignment_label,
    space_after,
    margin,
    title_size,
    heading1_size,
    heading2_size,
    heading3_size,
    add_page_numbers
):
    doc = Document(io.BytesIO(file_bytes))

    for section in doc.sections:
        section.top_margin = Inches(margin)
        section.bottom_margin = Inches(margin)
        section.left_margin = Inches(margin)
        section.right_margin = Inches(margin)

    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph_text(paragraph)
        if not text:
            continue

        final_type = final_type_map.get(idx, "Body")

        props = format_properties(
            final_type,
            body_size,
            title_size,
            heading1_size,
            heading2_size,
            heading3_size,
            body_alignment_label,
            line_spacing,
            space_after
        )

        if props["alignment"] == "Center":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif props["alignment"] == "Left":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif props["alignment"] == "Right":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif props["alignment"] == "Justified":
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        else:
            paragraph.alignment = alignment_enum

        paragraph.paragraph_format.space_before = Pt(props["space_before"])
        paragraph.paragraph_format.space_after = Pt(props["space_after"])
        paragraph.paragraph_format.line_spacing = props["line_spacing"]

        for run in paragraph.runs:
            set_run_font(
                run,
                font_name,
                props["size"],
                props["bold"] if final_type != "Body" else None
            )

    for table in doc.tables:
        format_table(table, font_name, body_size)

    if add_page_numbers:
        for section in doc.sections:
            footer = section.footer
            paragraph = (
                footer.paragraphs[0]
                if footer.paragraphs
                else footer.add_paragraph()
            )
            paragraph.clear()
            add_page_number(paragraph)

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# -------------------------------------------------------
# Session state
# -------------------------------------------------------

if "manual_overrides" not in st.session_state:
    st.session_state.manual_overrides = {}

if "last_file_sig" not in st.session_state:
    st.session_state.last_file_sig = None

# -------------------------------------------------------
# UI
# -------------------------------------------------------

uploaded_file = st.file_uploader(
    "Upload a Word document",
    type=["docx"],
    help="Only .docx files are supported."
)

if uploaded_file:
    file_bytes = uploaded_file.getvalue()
    current_sig = file_signature(file_bytes)

    if st.session_state.last_file_sig != current_sig:
        st.session_state.manual_overrides = {}
        st.session_state.last_file_sig = current_sig
        st.session_state.pop("formatted_doc", None)
        st.session_state.pop("output_name", None)

    st.success(f"Loaded: {uploaded_file.name}")

    left, right = st.columns(2)

    with left:
        st.subheader("1. Body formatting")

        font_name = st.selectbox(
            "Font",
            ["Times New Roman", "Arial", "Calibri", "Aptos", "Georgia"],
            index=0
        )

        body_size = st.number_input(
            "Body font size",
            min_value=8,
            max_value=20,
            value=12
        )

        line_spacing = st.selectbox(
            "Line spacing",
            [1.0, 1.15, 1.5, 2.0],
            index=2
        )

        alignment_label = st.selectbox(
            "Body alignment",
            ["Justified", "Left", "Center", "Right"],
            index=0
        )

        alignment_map = {
            "Justified": WD_ALIGN_PARAGRAPH.JUSTIFY,
            "Left": WD_ALIGN_PARAGRAPH.LEFT,
            "Center": WD_ALIGN_PARAGRAPH.CENTER,
            "Right": WD_ALIGN_PARAGRAPH.RIGHT
        }

        space_after = st.number_input(
            "Space after paragraph (pt)",
            min_value=0,
            max_value=30,
            value=6
        )

        margin = st.number_input(
            "Margins (inches)",
            min_value=0.3,
            max_value=2.0,
            value=1.0,
            step=0.1
        )

    with right:
        st.subheader("2. Structure & heading sizes")

        auto_detect = st.checkbox(
            "Auto-detect headings",
            value=True,
            help=(
                "ON: infer headings from numbering, capitalization, bold text and "
                "original font size. OFF: use only existing Word heading styles."
            )
        )

        title_size = st.number_input("Title size", 12, 36, 20)
        heading1_size = st.number_input("Heading 1 size", 10, 30, 16)
        heading2_size = st.number_input("Heading 2 size", 10, 26, 14)
        heading3_size = st.number_input("Heading 3 size", 10, 24, 13)

        add_page_numbers = st.checkbox("Add page numbers", value=True)

    st.divider()
    st.subheader("3. Review structure & proposed formatting")

    base_df = build_base_preview(
        file_bytes,
        auto_detect,
        body_size
    )

    base_df = apply_saved_overrides(
        base_df,
        st.session_state.manual_overrides
    )

    display_df = add_output_columns(
        base_df,
        body_size,
        title_size,
        heading1_size,
        heading2_size,
        heading3_size,
        alignment_label,
        line_spacing,
        space_after
    )

    if display_df.empty:
        st.info("No text paragraphs were found.")
        edited_df = display_df

    else:
        st.caption(
            "Change **Final type** when detection is wrong. The output-size, alignment "
            "and spacing columns show exactly what will be applied."
        )

        edited_df = st.data_editor(
            display_df,
            use_container_width=True,
            hide_index=True,
            num_rows="fixed",
            disabled=[
                "Paragraph index",
                "Detected as",
                "Text preview",
                "Original style",
                "Original size",
                "Originally bold",
                "Output size",
                "Output bold",
                "Output alignment",
                "Line spacing",
                "Space after"
            ],
            column_config={
                "Final type": st.column_config.SelectboxColumn(
                    "Final type",
                    options=CLASS_OPTIONS,
                    required=True
                ),
                "Text preview": st.column_config.TextColumn(
                    "Text preview",
                    width="large"
                ),
                "Output size": st.column_config.NumberColumn(
                    "Output size (pt)"
                ),
                "Space after": st.column_config.NumberColumn(
                    "Space after (pt)"
                )
            },
            key=f"editor_{current_sig}"
        )

        new_overrides = {}

        for _, row in edited_df.iterrows():
            idx = int(row["Paragraph index"])

            if row["Final type"] != row["Detected as"]:
                new_overrides[idx] = row["Final type"]

        st.session_state.manual_overrides = new_overrides

        # Recalculate output columns AFTER any manual edits.
        edited_df = add_output_columns(
            edited_df.drop(
                columns=[
                    "Output size",
                    "Output bold",
                    "Output alignment",
                    "Line spacing",
                    "Space after"
                ],
                errors="ignore"
            ),
            body_size,
            title_size,
            heading1_size,
            heading2_size,
            heading3_size,
            alignment_label,
            line_spacing,
            space_after
        )

        counts = edited_df["Final type"].value_counts()
        c1, c2, c3, c4, c5 = st.columns(5)
        c1.metric("Title", int(counts.get("Title", 0)))
        c2.metric("Heading 1", int(counts.get("Heading 1", 0)))
        c3.metric("Heading 2", int(counts.get("Heading 2", 0)))
        c4.metric("Heading 3", int(counts.get("Heading 3", 0)))
        c5.metric("Body", int(counts.get("Body", 0)))

        col_a, col_b = st.columns([1, 3])

        with col_a:
            if st.button("Reset manual corrections"):
                st.session_state.manual_overrides = {}
                st.rerun()

        with col_b:
            st.caption(
                f"Auto detection: {'ON' if auto_detect else 'OFF'} • "
                f"Manual corrections: {len(st.session_state.manual_overrides)}"
            )

    st.divider()
    st.subheader("4. Print-style preview")

    preview_enabled = st.checkbox(
        "Show formatted document preview",
        value=True,
        help="Shows an approximate print-style view before the Word file is created."
    )

    if preview_enabled and not edited_df.empty:
        final_type_map = build_final_type_map(edited_df)

        preview_html = render_print_preview(
            file_bytes=file_bytes,
            final_type_map=final_type_map,
            font_name=font_name,
            body_size=body_size,
            title_size=title_size,
            heading1_size=heading1_size,
            heading2_size=heading2_size,
            heading3_size=heading3_size,
            body_alignment=alignment_label,
            line_spacing=line_spacing,
            space_after=space_after,
            margin=margin
        )

        st.components.v1.html(
            preview_html,
            height=850,
            scrolling=True
        )

        st.caption(
            "This preview is for structural and visual review. Exact line wrapping, "
            "table layout and page breaks can differ slightly in Microsoft Word."
        )

    st.divider()
    st.subheader("5. Generate & download")

    if not edited_df.empty:
        final_type_map = build_final_type_map(edited_df)
    else:
        final_type_map = {}

    if st.button(
        "✨ Generate formatted Word document",
        type="primary",
        use_container_width=True
    ):
        try:
            result = format_document(
                file_bytes=file_bytes,
                final_type_map=final_type_map,
                font_name=font_name,
                body_size=body_size,
                line_spacing=line_spacing,
                alignment_enum=alignment_map[alignment_label],
                body_alignment_label=alignment_label,
                space_after=space_after,
                margin=margin,
                title_size=title_size,
                heading1_size=heading1_size,
                heading2_size=heading2_size,
                heading3_size=heading3_size,
                add_page_numbers=add_page_numbers
            )

            st.session_state.formatted_doc = result.getvalue()
            st.session_state.output_name = uploaded_file.name.replace(
                ".docx",
                "_formatted.docx"
            )

            st.success(
                "Document generated. Review the preview above, then download when ready."
            )

        except Exception as exc:
            st.error(f"Could not format the document: {exc}")

    if "formatted_doc" in st.session_state:
        st.download_button(
            "⬇️ Download formatted Word document",
            data=st.session_state.formatted_doc,
            file_name=st.session_state.output_name,
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )

else:
    st.info(
        "Upload a .docx file to review its structure and see a print-style preview."
    )

st.divider()
st.caption(
    "Manual structure corrections always take priority over automatic heading detection."
)
